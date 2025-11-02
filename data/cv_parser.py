# -*- coding: utf-8 -*-
"""
CV Parser with Mistral-First Strategy + Solid Fallbacks (English-only)
- Primary: LLM (Mistral via ollama) structured JSON with a strict schema
- Fallbacks: Regex/OCR extractors for personal info, education, languages, experience
- Output: DOCX in a reference-like layout (name/role rectangle + ordered sections)
- Supports many CV styles/roles; normalizes section names and experience entries

IMPORTANT: PERSONAL DATA renders ONLY these fields (and nothing else):
- Nationality
- Birth Date
- Gender
- Residence

Batch mode:
- CLI: python cv_parser.py <files and/or folders> --recursive --job-title "..."
- GUI fallback if no CLI inputs: pick multiple files or a whole folder
"""

import os
import re
import json
import argparse
from pathlib import Path
import docx
import tkinter as tk
from tkinter import simpledialog, filedialog, messagebox
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from typing import Dict, List, Optional, Any
import ollama

# ====== CONFIG ======
OUTPUT_DIR = r"DD:\cvprojfiles\outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

OCR_LANG = "eng"
SUPPORTED_EXTS = {".pdf", ".docx", ".doc"}

# ====== SECTION NORMALIZATION MAP ======
SECTION_ALIASES = {
    "WORK EXPERIENCE": "CHRONOLOGICAL EXPERIENCE RECORD",
    "EMPLOYMENT HISTORY": "CHRONOLOGICAL EXPERIENCE RECORD",
    "PROFESSIONAL EXPERIENCE": "CHRONOLOGICAL EXPERIENCE RECORD",
    "EXPERIENCE": "CHRONOLOGICAL EXPERIENCE RECORD",
    "WORK_EXPERIENCE": "CHRONOLOGICAL EXPERIENCE RECORD",
    "EXPERIENCE RECORD": "CHRONOLOGICAL EXPERIENCE RECORD",
    "TRAINING": "TRAINING COURSES AND CERTIFICATIONS",
    "COURSES": "TRAINING COURSES AND CERTIFICATIONS",
    "CERTIFICATIONS": "TRAINING COURSES AND CERTIFICATIONS",
    "QUALIFICATIONS": "EDUCATION",
    "ACADEMICS": "EDUCATION",
    "SKILLS": "COMPUTER SKILLS",
    "TECHNICAL SKILLS": "COMPUTER SKILLS",
    "IT SKILLS": "COMPUTER SKILLS",
    "COMPUTER KNOWLEDGE": "COMPUTER SKILLS",
    "LANGUAGE": "LANGUAGES",
    "LANGUAGE SKILLS": "LANGUAGES",
    "PERSONAL INFORMATION": "PERSONAL DATA",
    "PERSONAL DETAILS": "PERSONAL DATA",
}

# ====== UTIL ======
def normalize_sections(sections: Dict[str, Any]) -> Dict[str, Any]:
    normalized: Dict[str, Any] = {}
    for key, value in (sections or {}).items():
        key_clean = key.replace("_", " ").strip().upper()
        norm_key = SECTION_ALIASES.get(key_clean, key_clean)
        if norm_key not in normalized:
            normalized[norm_key] = []
        if isinstance(value, list):
            normalized[norm_key].extend(value)
        elif isinstance(value, dict):
            vals = []
            for _, v in sorted(value.items()):
                vals.append(v)
            normalized[norm_key].extend(vals)
        elif value:
            normalized[norm_key].append(value)
    return normalized

def safe_json_loads(raw: str):
    def clean_json_string(s: str) -> str:
        s = s.replace("…", "").replace("...", "")
        s = re.sub(r"//.*", "", s)
        s = re.sub(r"/\*.*?\*/", "", s, flags=re.S)
        s = re.sub(r",\s*([}\]])", r"\1", s)
        s = re.sub(r"([}\]])\s*([{\[])", r"\1,\2", s)
        s = re.sub(r"[\x00-\x1f\x7f]", "", s)
        return s.strip()
    attempt = raw
    for _ in range(6):
        try:
            return json.loads(attempt)
        except json.JSONDecodeError:
            attempt = clean_json_string(attempt)
    return None

def first_text(s: Optional[str]) -> str:
    return (s or "").strip()

def strip_trailing_punct(s: str) -> str:
    return re.sub(r"[,\.;:\-\s]+$", "", s or "").strip()

# ====== Robust name helpers ======
BRAND_TOKENS = {
    "enixo", "gia", "resume", "cv", "curriculum", "vitae", "profile",
    "summary", "objective", "experience", "education", "skills", "projects"
}

def is_brandish_line(s: str) -> bool:
    t = s.strip()
    low = t.lower()
    if re.search(r"https?://|www\.|@", low):
        return True
    if re.search(r"\d", low):
        return True
    if re.match(r"^[A-Za-z]{2,12}\s*-\s*[A-Za-z]{2,12}$", t):
        return True
    if any(tok in low for tok in BRAND_TOKENS):
        return True
    return False

def name_from_email(email: str) -> str:
    if not email:
        return ""
    m = re.match(r"([^@]+)@", email)
    if not m:
        return ""
    local = re.sub(r"[\d_]+", " ", m.group(1))
    parts = re.split(r"[.\-\s]+", local)
    parts = [p for p in parts if p and len(p) >= 2]
    cand = " ".join(p.capitalize() for p in parts[:5]).strip()
    return cand if len(cand.split()) >= 2 else ""

def guess_name_from_filename(path: str) -> str:
    base = os.path.splitext(os.path.basename(path))[0]
    base = re.sub(r"(?i)\b(cv|resume|final|updated|new|v\d+)\b", "", base)
    base = re.sub(r"[_+\-\(\)\[\]\.]", " ", base)
    base = re.sub(r"\s{2,}", " ", base).strip()
    if not base:
        return ""
    toks = [t for t in base.split() if re.match(r"^[A-Za-z][A-Za-z'.-]{1,}$", t)]
    toks = toks[:5]
    cand = " ".join(w.capitalize() for w in toks)
    return cand if 2 <= len(cand.split()) <= 5 else ""

def extract_name_safely(cv_text: str, email: str = "", filename: str = "") -> str:
    m = re.search(r"(?i)\bName\s*[:\-–]\s*([^\n\r]{3,80})", cv_text)
    if m:
        cand = re.sub(r"[,;|]+$", "", m.group(1).strip())
        if 2 <= len(cand.split()) <= 6 and not is_brandish_line(cand):
            return cand
    cand = name_from_email(email)
    if cand and not is_brandish_line(cand):
        return cand
    cand = guess_name_from_filename(filename)
    if cand and not is_brandish_line(cand):
        return cand
    lines = [l.strip() for l in cv_text.splitlines() if l.strip()]
    head = lines[:20]
    bad_starts = re.compile(r"(?i)^(personal|profile|objective|summary|contact|experience|work|education|skills|languages|training|courses|projects)\b")
    for l in head:
        if bad_starts.search(l):
            continue
        if is_brandish_line(l):
            continue
        if re.search(r"[@\d]", l):
            continue
        toks = l.split()
        if not (2 <= len(toks) <= 6):
            continue
        caplike = sum([1 for t in toks if re.match(r"^[A-Z][a-z'.-]+$|^[A-Z]{2,}$", t)])
        if caplike >= max(2, len(toks)-1):
            return re.sub(r"[,\.;:\-]+$", "", l).strip()
    return ""

# ====== OCR / TEXT EXTRACTION ======
def extract_text_from_pdf(pdf_path: str) -> List[str]:
    text_blocks: List[str] = []
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text = page.get_text().strip()
            if text:
                text_blocks.extend(line.strip() for line in text.split("\n") if line.strip())
            else:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                try:
                    ocr_text = pytesseract.image_to_string(img, lang=OCR_LANG)
                except Exception:
                    ocr_text = pytesseract.image_to_string(img)
                text_blocks.extend(line.strip() for line in ocr_text.split("\n") if line.strip())
    return text_blocks

def get_all_text(doc) -> List[str]:
    text_blocks: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_blocks.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_blocks.append(cell.text.strip())
    return text_blocks

# ====== LLM (Mistral-first) ======
MISTRAL_SCHEMA_EXAMPLE = """
{
  "name": "Full name",
  "headline": "Current job title or specialization",
  "contacts": {
    "Email": "email@domain.com",
    "Phones": ["+201234567890", "+971501234567"],
    "Address": "City, Country"
  },
  "personal_data": {
    "Nationality": "Nationality or empty",
    "Birth Date": "DD/MM/YYYY or YYYY-MM-DD or empty",
    "Gender": "Male/Female or empty",
    "Residence": "City, Country or empty"
  },
  "sections": {
    "EDUCATION": [
      "Bachelor in Mechanical Engineering, Cairo University, 2012",
      "Diploma in Industrial Maintenance, 2015"
    ],
    "LANGUAGES": [
      "Arabic: Native",
      "English: Very Good"
    ],
    "COMPUTER SKILLS": [
      "MS Office, AutoCAD, SAP PM"
    ],
    "TRAINING COURSES AND CERTIFICATIONS": [
      "OSHA 30h, NEBOSH IGC"
    ],
    "CHRONOLOGICAL EXPERIENCE RECORD": [
      {
        "Company": "Saudi Aramco",
        "Position": "Mechanical Fitter",
        "Duration": "FROM July 2018 TO Jan 2020",
        "Location": "Dhahran, KSA",
        "Tasks": [
          "Overhauling pumps and valves",
          "Preventive maintenance on rotating equipment"
        ]
      }
    ]
  }
}
""".strip()

def mistral_parse_cv(cv_text: str) -> Optional[Dict]:
    sys_prompt = (
        "You are an expert CV parser for English resumes. "
        "Return strictly valid JSON conforming to the schema. "
        "Do not include explanations. Ensure arrays are arrays and objects are objects. "
        "Infer reasonable durations/locations if clearly present near role lines. "
        "One language per item in LANGUAGES. Use consistent month abbreviations (e.g., Jan, Feb, Mar...)."
    )
    user_prompt = f"""
Schema (example — follow keys & shapes exactly, but fill with this CV content):
{MISTRAL_SCHEMA_EXAMPLE}

CV TEXT:
{cv_text}
""".strip()

    try:
        resp = ollama.chat(
            model="mistral",
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": user_prompt},
            ],
            options={"temperature": 0.1}
        )
        raw = resp["message"]["content"].strip()
        start, end = raw.find("{"), raw.rfind("}")
        if start == -1 or end == -1:
            return None
        cleaned = raw[start:end + 1]
        parsed = safe_json_loads(cleaned)
        if not parsed:
            return None

        if "sections" in parsed and isinstance(parsed["sections"], dict):
            parsed["sections"] = normalize_sections(parsed["sections"])

        contacts = parsed.get("contacts", {}) or {}
        phones = contacts.get("Phones", [])
        if isinstance(phones, str):
            phones = [p.strip() for p in re.split(r"[;/,]| or ", phones) if p.strip()]
            contacts["Phones"] = phones
        parsed["contacts"] = contacts

        exp = parsed.get("sections", {}).get("CHRONOLOGICAL EXPERIENCE RECORD")
        if isinstance(exp, dict):
            parsed["sections"]["CHRONOLOGICAL EXPERIENCE RECORD"] = list(exp.values())
        return parsed
    except Exception:
        return None

# ====== Heuristics for nationality & gender inference ======
COUNTRY_KEYWORDS = {
    "egypt": "Egyptian", "sohag": "Egyptian", "ismailia": "Egyptian", "alexandria": "Egyptian",
    "cairo": "Egyptian", "giza": "Egyptian", "mansoura": "Egyptian", "assiut": "Egyptian", "assuit": "Egyptian",
    "libya": "Libyan", "tripoli": "Libyan", "benghazi": "Libyan",
    "saudi": "Saudi", "saudi arabia": "Saudi", "riyadh": "Saudi", "jeddah": "Saudi",
    "uae": "Emirati", "dubai": "Emirati", "abu dhabi": "Emirati", "abudhabi": "Emirati",
    "qatar": "Qatari", "oman": "Omani", "kuwait": "Kuwaiti", "bahrain": "Bahraini",
    "sudan": "Sudanese", "morocco": "Moroccan", "tunisia": "Tunisian", "algeria": "Algerian",
    "india": "Indian", "pakistan": "Pakistani", "jordan": "Jordanian", "lebanon": "Lebanese",
    "syria": "Syrian", "iraq": "Iraqi", "yemen": "Yemeni", "turkey": "Turkish"
}

DIAL_CODE_TO_COUNTRY = {
    "+20": "Egyptian", "+218": "Libyan", "+966": "Saudi", "+971": "Emirati",
    "+974": "Qatari", "+968": "Omani", "+965": "Kuwaiti", "+973": "Bahraini",
    "+249": "Sudanese", "+212": "Moroccan", "+216": "Tunisian", "+213": "Algerian",
    "+91": "Indian", "+92": "Pakistani", "+962": "Jordanian", "+961": "Lebanese",
    "+963": "Syrian", "+964": "Iraqi", "+967": "Yemeni", "+90": "Turkish"
}

MALE_NAME_TOKENS = {
    "ahmed","mohamed","muhammad","abdul","abdel","abd","ali","hassan","hussein",
    "khaled","mahmoud","mostafa","moustafa","omar","youssef","yusuf","ibrahim",
    "reda","tarek","tareq","ramy","rami","hazem","mohammad","abdullah","abdalla","amin","shoman"
}
FEMALE_NAME_TOKENS = {
    "fatma","fatima","asmaa","asma","mariam","maryam","sara","sarah","aya",
    "eman","iman","doaa","dina","salma","noha","shimaa","hend","hoda","huda"
}

def infer_nationality_from_text(text: str) -> str:
    low = text.lower()
    for key, nat in COUNTRY_KEYWORDS.items():
        if key in low:
            return nat
    for code, nat in DIAL_CODE_TO_COUNTRY.items():
        if code in text:
            return nat
    return "Not specified"

def infer_gender_from_name_or_text(name: str, text: str) -> str:
    low_name = (name or "").lower()
    low_text = text.lower()
    if re.search(r"\b(mr\.|male)\b", low_text): return "Male"
    if re.search(r"\b(ms\.|mrs\.|female)\b", low_text): return "Female"
    tokens = re.findall(r"[a-z]+", low_name)
    if any(t in MALE_NAME_TOKENS for t in tokens): return "Male"
    if any(t in FEMALE_NAME_TOKENS for t in tokens): return "Female"
    return "Not specified"

# ====== FALLBACK EXTRACTORS ======
def normalize_language_level(level: str) -> str:
    if not level:
        return "Not specified"
    level = level.lower()
    if "mother" in level or "native" in level: return "Native"
    if "fluent" in level: return "Fluent"
    if "excellent" in level or "very good" in level: return "Very Good"
    if "good" in level and "very" not in level: return "Good"
    if "fair" in level or "basic" in level: return "Fair"
    return "Not specified"

DATE_PATTERNS = [
    r"([0-3]?\d[/-][01]?\d[/-](?:19|20)\d{2})",
    r"((?:19|20)\d{2}-[01]\d-[0-3]\d)",
    r"([0-3]?\d\s+[A-Za-z]{3,9}\s+(?:19|20)\d{2})",
    r"([A-Za-z]{3,9}\s+[0-3]?\d,\s*(?:19|20)\d{2})",
]

LABELS = {
    "nationality": r"(?:Nationality)",
    "gender": r"(?:Gender)",
    "male": r"(?:Male)",
    "female": r"(?:Female)",
    "birth_date": r"(?:Date of Birth|Birth Date)",
    "address": r"(?:Address|Residence)",
    "position": r"(?:Position|Job title|Title|Designation|Role)",
    "company": r"(?:Company|Employer)",
    "role_desc": r"(?:Role description|Responsibilities|Duties|Tasks)",
}

def find_first_date(text: str) -> str:
    for pat in DATE_PATTERNS:
        m = re.search(pat, text, re.I)
        if m:
            return m.group(1).replace("-", "/")
    return ""

def extract_personal_data_fallback(cv_text: str, current_data: dict) -> dict:
    data = (current_data or {}).copy()
    text = cv_text

    if not first_text(data.get("Birth Date")):
        if re.search(LABELS["birth_date"] + r"\s*[:\-–]?\s*([^\n\r]+)", text, re.I):
            src = re.search(LABELS["birth_date"] + r"\s*[:\-–]?\s*([^\n\r]+)", text, re.I).group(1)
            bd = find_first_date(src)
        else:
            bd = find_first_date(text)
        data["Birth Date"] = bd or "Not specified"

    if not first_text(data.get("Residence")):
        m = re.search(LABELS["address"] + r"\s*[:\-–]?\s*([^\n\r]+)", text, re.I)
        if m:
            line = strip_trailing_punct(m.group(1).strip())
            lines = cv_text.splitlines()
            idx = None
            for i, l in enumerate(lines):
                if m.group(0) in l or m.group(1) in l:
                    idx = i
                    break
            if idx is not None and (m.group(1).strip().endswith(",") or lines[idx].strip().endswith(",")):
                if idx + 1 < len(lines):
                    nxt = strip_trailing_punct(lines[idx + 1].strip())
                    if nxt and len(nxt.split()) >= 2:
                        line = f"{line}, {nxt}"
            data["Residence"] = line
        else:
            data.setdefault("Residence", "Not specified")

    if not first_text(data.get("Nationality")):
        m = re.search(LABELS["nationality"] + r"\s*[:\-–]?\s*([A-Za-z ]{3,40})", text, re.I)
        if m:
            data["Nationality"] = strip_trailing_punct(m.group(1))
        else:
            cand = infer_nationality_from_text(" ".join([text, data.get("Residence","")]))
            data["Nationality"] = cand

    if not first_text(data.get("Gender")):
        m = re.search(LABELS["gender"] + r"\s*[:\-–]?\s*(" + LABELS["male"] + r"|" + LABELS["female"] + r")", text, re.I)
        if m:
            g = m.group(1).lower()
            data["Gender"] = "Male" if re.search(LABELS["male"], g, re.I) else "Female"
        else:
            possible_name = re.search(r"(?i)\bName\s*:\s*([^\n\r]+)", text)
            name_val = possible_name.group(1).strip() if possible_name else ""
            data["Gender"] = infer_gender_from_name_or_text(name_val, text)

    for field in ["Nationality", "Birth Date", "Gender", "Residence"]:
        if not first_text(data.get(field)):
            data[field] = "Not specified"
    return data

def extract_personal_info_from_text(cv_text: str) -> dict:
    info = {
        "Name": "",
        "Headline": "",
        "Company": "",
        "Email": "",
        "Phones": [],
        "Address": "",
        "Birth Date": "",
        "Years of Service": "",
    }
    text = re.sub(r"[ \t]+", " ", cv_text)
    text = re.sub(r"\s*:\s*", ": ", text)

    m = re.search(r"(?i)\bName\s*:\s*([^\n\r]+)", text)
    if m: info["Name"] = m.group(1).strip()

    m = re.search(r"(?i)\b(Current\s*(Position|Role|Title)|Headline|Designation|Title)\s*:\s*([^\n\r]+)", text)
    if m: info["Headline"] = m.group(3).strip()

    m = re.search(LABELS["company"] + r"\s*:\s*([^\n\r]+)", text, re.I)
    if m: info["Company"] = m.group(1).strip()

    m = re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", text, re.I)
    if m: info["Email"] = m.group(0)

    phones = re.findall(r"(?:\+?\d{1,3}[-\s]?)?(?:\d{2,4}[-\s]?){2,}\d{2,}", text)
    clean_phones = []
    for p in phones:
        p2 = re.sub(r"[^\d+]", "", p)
        if len(re.sub(r"\D", "", p2)) >= 9 and p2 not in clean_phones:
            clean_phones.append(p2)
    info["Phones"] = clean_phones[:3]

    m = re.search(LABELS["address"] + r"\s*:\s*([^\n\r]+)", text, re.I)
    if m:
        addr = strip_trailing_punct(m.group(1))
        info["Address"] = addr

    bd = find_first_date(text)
    if bd: info["Birth Date"] = bd

    m = re.search(r"(?i)\bYears of Service\s*:\s*([^\n\r]+)", text)
    if m: info["Years of Service"] = m.group(1).strip()

    return info

def extract_education_from_text(cv_text: str) -> List[str]:
    lines = [ln.strip() for ln in cv_text.splitlines() if ln.strip()]
    results, seen = [], set()
    deg_tokens = r"(bachelor|b\.?\s?sc|bsc|master|m\.?\s?sc|ph\.?d|diploma|degree)"
    inst_tokens = r"(university|faculty|institute|academy|college)"
    year = r"(?:19|20)\d{2}"
    blacklist = re.compile(r"(autocad|solidworks|mastercam|artcam|edgecam|icdl|ms office|sap|excel|word|powerpoint)", re.I)

    for i, ln in enumerate(lines):
        if blacklist.search(ln):
            continue
        deg_hit = re.search(deg_tokens, ln, re.I)
        inst_hit = re.search(inst_tokens, ln, re.I)
        if deg_hit or inst_hit:
            context = " ".join([l for l in lines[i:i+3] if not blacklist.search(l)])
            if not re.search(deg_tokens, context, re.I) and not re.search(r"graduat", context, re.I):
                continue
            m_deg = re.search(r"(?i)(Bachelor|B\.?\s?Sc\.?|Master|M\.?\s?Sc\.?|Ph\.?D\.?|Diploma)([^,.;]{0,60})", context)
            deg = m_deg.group(0).strip(" -:.") if m_deg else ""
            m_inst = re.search(r"(?i)([A-Za-z ]*\bUniversity\b(?: of [A-Za-z ]+)?)", context)
            inst = (m_inst.group(1).strip() if m_inst else "")
            m_years = re.findall(year, context)
            y = m_years[-1] if m_years else ""
            m_field = re.search(r"(?i)\b(?:in)\s+([A-Za-z &]{2,50})", context)
            field = m_field.group(1).strip() if m_field else ""
            parts = [p for p in [deg, f"in {field}" if field else "", inst, y] if p]
            s = ", ".join(parts)
            s = re.sub(r"\s{2,}", " ", s).replace(" ,", ",").strip(", ")
            if s and s.lower() not in seen:
                seen.add(s.lower())
                results.append(s)
    return results

def extract_languages_from_text(cv_text: str) -> List[str]:
    langs_map = {
        "arabic": "Arabic", "english": "English", "french": "French", "german": "German",
        "italian": "Italian", "spanish": "Spanish", "turkish": "Turkish", "hindi": "Hindi"
    }
    levels_map = {
        "mother tongue": "Native", "native": "Native",
        "fluent": "Fluent",
        "excellent": "Very Good", "very good": "Very Good",
        "good": "Good",
        "fair": "Fair", "basic": "Basic",
        "reading and writing": "Good",
    }
    lines = [l.strip() for l in cv_text.splitlines() if l.strip()]
    rank = ["Not specified","Basic","Fair","Good","Very Good","Fluent","Native"]

    found = {}
    for ln in lines:
        low = ln.lower()
        for k, L in langs_map.items():
            if k in low:
                level = "Not specified"
                for key, val in levels_map.items():
                    if key in low:
                        level = val
                        break
                prev = found.get(L, "Not specified")
                if rank.index(level) > rank.index(prev):
                    found[L] = level

    if not found and re.search(r"(?i)Arabic.*mother tongue|mother tongue.*Arabic", cv_text):
        found["Arabic"] = "Native"

    return [f"{lang}: {lvl}" for lang, lvl in found.items()]

# ====== Experience helpers ======
MONTHS = {
    'jan':'Jan','january':'Jan','feb':'Feb','february':'Feb','mar':'Mar','march':'Mar',
    'apr':'Apr','april':'Apr','may':'May','jun':'Jun','june':'Jun','jul':'Jul','july':'Jul',
    'aug':'Aug','august':'Aug','sep':'Sep','sept':'Sep','september':'Sep',
    'oct':'Oct','october':'Oct','nov':'Nov','november':'Nov','dec':'Dec','december':'Dec'
}

def norm_month_word(s: str) -> str:
    low = s.lower().strip(".")
    return MONTHS.get(low, s[:3].title())

def extract_experience_from_text(cv_text: str) -> List[dict]:
    text = re.sub(r"[ \t]+", " ", cv_text)
    text = text.replace("\r", "")
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    date_re_list = [
        re.compile(r"(?i)\bfrom\s+([A-Za-z]{3,9}\.?\s+\d{4})\s+(?:to|till)\s+([A-Za-z]{3,9}\.?\s+\d{4}|now|present|until now|current)"),
        re.compile(r"(?i)\bfrom\s+([01]?\d)[/.-]((?:19|20)\d{2})\s+(?:to|till)\s+([01]?\d)[/.-]((?:19|20)\d{2}|now|present|until now|current)"),
        re.compile(r"(?i)\bfrom\s+((?:19|20)\d{1,2})\s+(?:to|till)\s+((?:19|20)\d{2}|now|present|until now|current)"),
    ]

    def _mon_to_word(num_str: str) -> str:
        try:
            n = int(num_str)
            names = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
            if 1 <= n <= 12:
                return names[n-1]
        except Exception:
            pass
        return num_str

    roles: List[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        dm = None
        for r in date_re_list:
            dm = r.search(line)
            if not dm and i+1 < len(lines):
                dm = r.search(lines[i+1])
            if dm:
                break

        if dm:
            def fmt_range(m):
                if m.re == date_re_list[0]:
                    start = re.split(r"\s+", m.group(1))
                    end = re.split(r"\s+", m.group(2))
                    s = f"{norm_month_word(start[0])} {start[-1]}"
                    e = m.group(2).strip()
                    if len(end) >= 2:
                        e = f"{norm_month_word(end[0])} {end[-1]}"
                    return f"FROM {s} TO {e.title()}"
                elif m.re == date_re_list[1]:
                    s_mon, s_year, e_mon, e_year = m.group(1), m.group(2), m.group(3), m.group(4)
                    s = f"{norm_month_word(_mon_to_word(s_mon))} {s_year}"
                    e = "Now" if re.match(r"(?i)now|present|until now|current", str(e_year)) else f"{norm_month_word(_mon_to_word(e_mon))} {e_year}"
                    return f"FROM {s} TO {e}"
                else:
                    s_year, e_year = m.group(1), m.group(2)
                    e = "Now" if re.match(r"(?i)now|present|until now|current", str(e_year)) else e_year
                    return f"FROM {s_year} TO {e}"

            dur = fmt_range(dm)

            context = " ".join(lines[max(0, i-4): i+12])
            m_comp = re.search(LABELS["company"] + r"\s*[: ]\s*([A-Za-z0-9 .,&\-()/]+)", context, re.I)
            company = strip_trailing_punct(m_comp.group(1)) if m_comp else ""

            m_pos = re.search(LABELS["position"] + r"\s*[: ]\s*([A-Za-z0-9 .,&\-()/]+)", context, re.I)
            position = strip_trailing_punct(m_pos.group(1)) if m_pos else ""
            if not position:
                m_pos2 = re.search(r"(?i)\b(As\s+)?([A-Z][A-Za-z ]{2,60})(?: at | with |,|\.|$)", context)
                if m_pos2:
                    position = strip_trailing_punct(m_pos2.group(2))

            task_lines = []
            for j in range(i, min(i+18, len(lines))):
                l = lines[j]
                if re.match(r"^\s*(?:[-•]|\d+\))", l) or re.search(r"(?i)\bmainten|\binstall|\bcommission|\boverhaul|\bvalve|\bpump|\bturbine|\bproject|\binspection|\boperation|\btesting|\bstart[- ]?up", l):
                    task_lines.append(re.sub(r"^\s*(?:[-•]|\d+\)|\d+\-)\s*", "", l).strip())

            clean = []
            seen_task = set()
            for t in task_lines:
                t = re.sub(r"\s{2,}", " ", t)
                if not t or len(t) < 3 or len(t) > 250:
                    continue
                if re.search(r"(?i)^(company|employer|position|job title|role description)\s*:", t):
                    continue
                if t.startswith(""):
                    t = t.replace("", "").strip()
                if t.lower().startswith("job title"):
                    continue
                if "FROM " in t:
                    continue
                t = strip_trailing_punct(t)
                if t and t.lower() not in seen_task:
                    seen_task.add(t.lower())
                    clean.append(t)

            if company or position or clean or dur:
                roles.append({
                    "Company": company,
                    "Position": position or "Role not specified",
                    "Duration": dur,
                    "Tasks": clean[:10]
                })
            i += 1
        i += 1

    dedup, seen = [], set()
    for r in roles:
        key = (r.get("Company","").lower(), r.get("Position","").lower(), r.get("Duration","").lower())
        if key not in seen:
            seen.add(key)
            dedup.append(r)
    return dedup

# ====== DOCX OUTPUT (STYLING-ONLY CHANGES) ======

def set_page_margins(doc, top=0.69, bottom=0.87, left=0.87, right=0.87):
    """Apply requested page margins in inches."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def _twips(inches: float) -> int:
    # 1 inch = 1440 twips
    return int(round(inches * 1440))

def add_name_job_rectangle(doc, name: str, job_title: str):
    """
    Header box:
    - Right-aligned
    - Stretched to full content width (between margins)
    - Thin black border
    - Tight internal padding and paragraph spacing
    """
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    table.autofit = False

    # Compute content width = page width - left/right margins
    section = doc.sections[0]
    content_width_in = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    # set python-docx column width (inches)
    try:
        table.columns[0].width = Inches(max(content_width_in - 0.02, 1.0))  # tiny safety offset
    except Exception:
        pass

    # Ensure tblPr exists
    tbl = table._element
    if tbl.tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    else:
        tblPr = tbl.tblPr

    # Explicit table width (so Word keeps it wide)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(_twips(content_width_in)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Borders
    tblBorders = OxmlElement('w:tblBorders')
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')        # ~1 pt border
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # Tight cell padding
    cellMar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        mar = OxmlElement(f'w:{side}')
        mar.set(qn('w:w'), '40')       # 40 twips ≈ 0.028"
        mar.set(qn('w:type'), 'dxa')
        cellMar.append(mar)
    tblPr.append(cellMar)

    # Row 1: Name (bold, right-aligned)
    cell_name = table.cell(0, 0)
    p_name = cell_name.paragraphs[0]
    p_name.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(0)
    run_name = p_name.add_run(first_text(name))
    run_name.font.name = 'Arial'
    run_name.font.size = Pt(18)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(0, 0, 0)

    # Row 2: Job title (right-aligned, tight spacing)
    cell_job = table.cell(1, 0)
    p_job = cell_job.paragraphs[0]
    p_job.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_job.paragraph_format.space_before = Pt(0)
    p_job.paragraph_format.space_after = Pt(0)
    run_job = p_job.add_run(first_text(job_title))
    run_job.font.name = 'Arial'
    run_job.font.size = Pt(12)
    run_job.font.bold = False
    run_job.font.color.rgb = RGBColor(0, 0, 0)

def add_section_heading(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name = 'Arial'
    run.font.size = Pt(16)  # 16pt per requirement
    run.font.bold = True
    return p

# --- employer/project formatting helpers
_ROLE_WORDS = re.compile(r"(?i)\b(engineer|technician|supervisor|manager|inspector|consultant|specialist|lead|senior|qa/qc|commission(?:ing)?|construction|site|hook[-\s]?up|e&?i|electrical|mechanical)\b")

def _clean_employer_only(company_raw: str) -> str:
    """keep company name only even if role-like text is mixed in."""
    s = first_text(company_raw)
    if not s:
        return s
    for kw in [" at ", " for ", " with "]:
        if kw in s.lower():
            left, right = re.split(kw, s, flags=re.I, maxsplit=1)
            if _ROLE_WORDS.search(left):
                s = right
    if "," in s:
        parts = [p.strip() for p in s.split(",")]
        if parts and _ROLE_WORDS.search(parts[0]) and len(parts) >= 2:
            s = parts[-1]
    return re.sub(r"\s{2,}", " ", s).strip(" .-")

def write_personal_block(doc, personal_ordered):
    add_section_heading(doc, "PERSONAL DATA")
    table = doc.add_table(rows=0, cols=3)
    table.autofit = False
    table.columns[0].width = Inches(1.3)
    table.columns[1].width = Inches(0.15)
    table.columns[2].width = Inches(4.5)

    for label, value in personal_ordered:
        val = first_text(value)
        if not val:
            continue
        row = table.add_row().cells
        p0 = row[0].paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.name = 'Arial'
        r0.font.size = Pt(10)
        r0.bold = False

        p1 = row[1].paragraphs[0]
        r1 = p1.add_run(":")
        r1.font.name = 'Arial'
        r1.font.size = Pt(10)

        p2 = row[2].paragraphs[0]
        r2 = p2.add_run(val)
        r2.font.name = 'Arial'
        r2.font.size = Pt(10)

    doc.add_paragraph("")

def write_list_block(doc, title: str, items: List[str]):
    if not items:
        return
    add_section_heading(doc, title)
    for s in items:
        s = first_text(s)
        if s and s.lower() != "not specified":
            para = doc.add_paragraph()
            run = para.add_run("• " + s)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
    doc.add_paragraph("")

def write_experience_block(doc, roles: List[dict]):
    if not roles:
        return
    add_section_heading(doc, "CHRONOLOGICAL EXPERIENCE RECORD")

    def add_kv(label: str, value: str):
        v = first_text(value)
        if not v or v.lower() == "not specified":
            return
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label} : ")
        r1.font.name = 'Arial'; r1.font.size = Pt(10); r1.bold = True
        r2 = p.add_run(v)
        r2.font.name = 'Arial'; r2.font.size = Pt(10)

    def clean(items):
        out, seen = [], set()
        for x in items or []:
            s = first_text(x).rstrip(" ,.;:-")
            if not s:
                continue
            low = s.lower()
            if low not in seen and low != "not specified":
                seen.add(low)
                out.append(s)
        return out

    proj_name_pat = re.compile(r"(?i)\bProject\s*(?:Name|Title)?\s*[:\-]\s*(.+)")

    def split_project(name_or_line: str) -> (str, str):
        t = first_text(name_or_line)
        if not t:
            return "", ""
        if "(" in t and ")" in t:
            pre = t.split("(", 1)[0].strip(" -,:;")
            par = t.split("(", 1)[1].rsplit(")", 1)[0].strip()
            if pre:
                return pre, par
        if " - " in t:
            parts = [p.strip(" -,:;") for p in t.split(" - ") if p.strip()]
            if parts:
                return parts[0], " - ".join(parts[1:]) if len(parts) > 1 else ""
        if "," in t:
            a, b = t.split(",", 1)
            return a.strip(" -,:;"), b.strip()
        return t, ""

    for r in roles:
        duration = first_text(r.get("Duration"))
        company  = _clean_employer_only(first_text(r.get("Company")))
        position = first_text(r.get("Position"))
        tasks    = clean(r.get("Tasks") or [])

        project_name, project_brief = "", ""
        new_tasks = []
        for t in tasks:
            m = proj_name_pat.search(t)
            if m and not project_name:
                pn, pb = split_project(m.group(1))
                project_name, project_brief = pn, pb
                continue
            new_tasks.append(t)
        tasks = new_tasks

        if not project_name and tasks:
            candidate_name, candidate_brief = split_project(tasks[0])
            used_first = False
            if len(candidate_name.split()) >= 2:
                project_name = candidate_name
                project_brief = project_brief or candidate_brief
                used_first = True
            elif not project_brief:
                project_brief = tasks[0]
                used_first = True
            if used_first:
                tasks = tasks[1:]

        if project_name and project_brief:
            project_line = f"{project_name} – {project_brief}"
        elif project_name:
            project_line = project_name
        else:
            project_line = project_brief

        add_kv("Dates", duration)
        add_kv("Employer", company)
        if project_line:
            add_kv("Project", project_line)
        add_kv("Job title", position)

        if tasks:
            p = doc.add_paragraph()
            rlbl = p.add_run("Job Description :")
            rlbl.font.name = 'Arial'; rlbl.font.size = Pt(10); rlbl.bold = True
            for t in tasks:
                para = doc.add_paragraph()
                run = para.add_run(f"- {t}")
                run.font.name = 'Arial'; run.font.size = Pt(10)

        doc.add_paragraph("")

def write_field_of_experience_block(doc, items: List[str]):
    items = [first_text(x) for x in (items or []) if first_text(x)]
    if not items:
        return
    p = doc.add_paragraph()
    r = p.add_run("Field of experience :")
    r.font.name = 'Arial'; r.font.size = Pt(10); r.bold = True
    for it in items:
        para = doc.add_paragraph()
        run = para.add_run(f"• {it}")
        run.font.name = 'Arial'; run.font.size = Pt(10)
    doc.add_paragraph("")

# ====== MERGE / CONSOLIDATE ======
def consolidate_structures(llm_out: Dict[str, Any], cv_text: str) -> Dict[str, Any]:
    llm_out = llm_out or {}
    sections = normalize_sections(llm_out.get("sections", {}))

    strong_pi = extract_personal_info_from_text(cv_text)
    contacts = llm_out.get("contacts", {}) or {}
    if isinstance(contacts.get("Phones"), str):
        contacts["Phones"] = [p.strip() for p in re.split(r"[;/,]| or ", contacts["Phones"]) if p.strip()]
    contacts.setdefault("Email", strong_pi.get("Email", ""))
    if not contacts.get("Phones"):
        contacts["Phones"] = strong_pi.get("Phones", [])
    contacts.setdefault("Address", strong_pi.get("Address", ""))

    personal_data = llm_out.get("personal_data", {}) or {}
    personal_data = extract_personal_data_fallback(cv_text, personal_data)
    if not personal_data.get("Birth Date"):
        personal_data["Birth Date"] = strong_pi.get("Birth Date", "") or "Not specified"
    if not personal_data.get("Residence"):
        personal_data["Residence"] = strip_trailing_punct(strong_pi.get("Address", "")) or "Not specified"

    name = first_text(llm_out.get("name")) or strong_pi.get("Name", "")
    headline = first_text(llm_out.get("headline")) or strong_pi.get("Headline", "")

    edu = sections.get("EDUCATION") or extract_education_from_text(cv_text)
    else_edu = [e for e in edu if "No Languages" not in str(e)]
    edu = else_edu

    langs = sections.get("LANGUAGES") or extract_languages_from_text(cv_text)

    exp = sections.get("CHRONOLOGICAL EXPERIENCE RECORD")
    if not exp:
        exp = extract_experience_from_text(cv_text)
    else:
        if isinstance(exp, dict):
            exp = list(exp.values())
        fixed = []
        for r in exp:
            if not isinstance(r, dict):
                continue
            fixed.append({
                "Company": first_text(r.get("Company")),
                "Position": first_text(r.get("Position")) or "Role not specified",
                "Duration": first_text(r.get("Duration")),
                "Location": first_text(r.get("Location")),
                "Tasks": [first_text(t) for t in (r.get("Tasks") or []) if first_text(t)]
            })
        exp = fixed

    comp = sections.get("COMPUTER SKILLS") or []
    train = sections.get("TRAINING COURSES AND CERTIFICATIONS") or sections.get("TRAINING") or []

    extra_sections = {k: v for k, v in sections.items() if k not in {
        "EDUCATION", "LANGUAGES", "COMPUTER SKILLS",
        "TRAINING COURSES AND CERTIFICATIONS", "CHRONOLOGICAL EXPERIENCE RECORD"
    }}

    return {
        "name": name,
        "headline": headline,
        "contacts": contacts,
        "personal_data": personal_data,
        "sections": {
            "EDUCATION": edu,
            "LANGUAGES": langs,
            "COMPUTER SKILLS": comp,
            "TRAINING COURSES AND CERTIFICATIONS": train,
            "CHRONOLOGICAL EXPERIENCE RECORD": exp,
            **extra_sections
        }
    }

# ====== CORE PROCESSOR ======
def process_cv(file_path: str, job_title: str = ""):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text_blocks = extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        try:
            doc_obj = docx.Document(file_path)
            text_blocks = get_all_text(doc_obj)
        except Exception:
            text_blocks = []
    else:
        print(f"❌ Unsupported file type: {ext}")
        return

    if not text_blocks and ext != ".pdf":
        print("⚠️ No text extracted; if this is a legacy .doc, consider saving as .docx or PDF and retry.")
    cv_text = "\n".join(text_blocks)

    parsed = mistral_parse_cv(cv_text)

    merged = consolidate_structures(parsed or {}, cv_text)

    strong_pi = extract_personal_info_from_text(cv_text)
    safe_name = extract_name_safely(cv_text, email=strong_pi.get("Email", ""), filename=file_path)
    if safe_name and (not merged.get("name") or is_brandish_line(merged.get("name", ""))):
        merged["name"] = safe_name

    pdata = merged.get("personal_data", {}) or {}
    if not first_text(pdata.get("Gender")) or pdata.get("Gender") == "Not specified":
        pdata["Gender"] = infer_gender_from_name_or_text(merged.get("name",""), cv_text)
        merged["personal_data"] = pdata

    # Create doc and apply margins
    doc_out = docx.Document()
    set_page_margins(doc_out, top=0.69, bottom=0.87, left=0.87, right=0.87)

    # Header box (styling only)
    add_name_job_rectangle(
        doc_out,
        merged.get("name", ""),
        merged.get("headline", "") or job_title
    )
    doc_out.add_paragraph("")

    pdata = merged.get("personal_data", {}) or {}
    contacts = merged.get("contacts", {}) or {}
    personal_ordered = [
        ("Nationality", pdata.get("Nationality","")),
        ("Birth Date", pdata.get("Birth Date","")),
        ("Gender", pdata.get("Gender","")),
        ("Residence", pdata.get("Residence","") or contacts.get("Address","")),
    ]
    write_personal_block(doc_out, personal_ordered)

    sections = merged.get("sections", {})
    write_list_block(doc_out, "EDUCATION", sections.get("EDUCATION") or [])
    write_list_block(doc_out, "LANGUAGES", sections.get("LANGUAGES") or [])
    write_list_block(doc_out, "COMPUTER SKILLS", sections.get("COMPUTER SKILLS") or [])
    write_list_block(doc_out, "TRAINING COURSES AND CERTIFICATIONS", sections.get("TRAINING COURSES AND CERTIFICATIONS") or [])
    write_experience_block(doc_out, sections.get("CHRONOLOGICAL EXPERIENCE RECORD") or [])
    write_field_of_experience_block(doc_out, sections.get("FIELD OF EXPERIENCE"))

    core = {"EDUCATION","LANGUAGES","COMPUTER SKILLS","TRAINING COURSES AND CERTIFICATIONS","CHRONOLOGICAL EXPERIENCE RECORD","FIELD OF EXPERIENCE"}
    for sec_name, items in sections.items():
        if sec_name in core:
            continue
        if isinstance(items, list) and items:
            write_list_block(doc_out, sec_name, [str(x) for x in items if first_text(str(x))])
        elif isinstance(items, dict) and items:
            flat = []
            for _, v in items.items():
                if isinstance(v, str) and first_text(v):
                    flat.append(v)
                elif isinstance(v, list):
                    flat.extend([str(x) for x in v if first_text(str(x))])
                elif isinstance(v, dict):
                    for k2, v2 in v.items():
                        if first_text(str(v2)):
                            flat.append(f"{k2}: {v2}")
            write_list_block(doc_out, sec_name, flat)

    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_path = os.path.join(OUTPUT_DIR, f"{base_name}_parsed.docx")
    counter = 1
    while os.path.exists(output_path):
        output_path = os.path.join(OUTPUT_DIR, f"{base_name}_parsed_{counter}.docx")
        counter += 1
    doc_out.save(output_path)
    print(f"✅ Saved parsed CV to: {output_path}")

# ====== BATCH HELPERS ======
def discover_input_files(inputs, recursive: bool) -> List[str]:
    files: List[str] = []
    for item in inputs:
        p = Path(item)
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS:
            files.append(str(p))
        elif p.is_dir():
            if recursive:
                for f in p.rglob("*"):
                    if f.suffix.lower() in SUPPORTED_EXTS:
                        files.append(str(f))
            else:
                for f in p.glob("*"):
                    if f.suffix.lower() in SUPPORTED_EXTS:
                        files.append(str(f))
    seen = set()
    uniq = []
    for f in files:
        if f not in seen:
            seen.add(f)
            uniq.append(f)
    return uniq

def run_batch(files: List[str], job_title: str = ""):
    if not files:
        print("⚠️ No input files found.")
        return
    ok, fail = 0, 0
    for fp in files:
        try:
            print(f"▶ Processing: {fp}")
            process_cv(fp, job_title=job_title)
            ok += 1
        except Exception as e:
            print(f"❌ Failed to process {fp}: {e}")
            fail += 1
    print(f"\n=== Batch finished ===\n✅ Success: {ok}\n❌ Failed: {fail}")

# ====== LAUNCHER ======
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Batch CV Parser (files and/or folders). If no inputs are provided, a GUI picker will open."
    )
    parser.add_argument("inputs", nargs="*", help="Files or folders to process")
    parser.add_argument("--recursive", "-r", action="store_true", help="Recurse into folders")
    parser.add_argument("--job-title", default="", help="Optional job title (overrides detected headline)")
    args = parser.parse_args()

    if args.inputs:
        files = discover_input_files(args.inputs, recursive=args.recursive)
        run_batch(files, job_title=args.job_title)
    else:
        root = tk.Tk()
        root.withdraw()

        choice = simpledialog.askstring(
            "Input mode",
            "Type 'files' to pick multiple files, or 'folder' to pick a whole folder:\n(default = files)",
            initialvalue="files"
        )
        choice = (choice or "files").strip().lower()

        if choice == "folder":
            folder = filedialog.askdirectory(title="Select a folder that contains CVs")
            if not folder:
                print("No folder selected.")
            else:
                rec = messagebox.askyesno("Recursive?", "Scan subfolders too?")
                job_title = simpledialog.askstring("Job Title (optional)", "Enter Job Title (leave blank to auto-detect):", initialvalue="")
                files = discover_input_files([folder], recursive=rec)
                run_batch(files, job_title=(job_title or ""))
        else:
            paths = filedialog.askopenfilenames(
                title="Select one or more CV files",
                filetypes=[("Supported Files", "*.docx *.doc *.pdf"),
                           ("Word Documents", "*.docx *.doc"),
                           ("PDF Files", "*.pdf")]
            )
            if not paths:
                print("No files selected.")
            else:
                job_title = simpledialog.askstring("Job Title (optional)", "Enter Job Title (leave blank to auto-detect):", initialvalue="")
                run_batch(list(paths), job_title=(job_title or ""))
